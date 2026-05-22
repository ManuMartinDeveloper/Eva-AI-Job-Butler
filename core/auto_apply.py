import os
import sys
import time
import json
import re
import argparse
import docx
from datetime import datetime
from playwright.sync_api import sync_playwright
from playwright_stealth import stealth_sync

# --- Setup Paths ---
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(_SCRIPT_DIR)
sys.path.append(_PROJECT_ROOT)

from core.db import SessionLocal, ProfileFact, Job, Application, AgentLog
from core.profile_memory_resume_phase2 import get_llm, get_qdrant_client, get_rag_context, parse_json_output

def get_user_session_path(user_id: int) -> str:
    path = os.path.join(_PROJECT_ROOT, "outputs", str(user_id), "agent_session.json")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return path

def get_user_screenshot_path(user_id: int) -> str:
    path = os.path.join(_PROJECT_ROOT, "outputs", str(user_id), "agent_screenshot.png")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return path

def update_session(user_id: int, status, logs, screenshot_path=None, command=None, job_info=None):
    """Updates the shared session state file for FastAPI to read."""
    session_path = get_user_session_path(user_id)
    session_data = {}
    if os.path.exists(session_path):
        try:
            with open(session_path, 'r') as f:
                session_data = json.load(f)
        except:
            pass

    session_data["status"] = status
    session_data["logs"] = logs
    if screenshot_path:
        session_data["screenshot_path"] = f"outputs/{user_id}/agent_screenshot.png"
    if command is not None:
        session_data["command"] = command
    if job_info:
        session_data.update(job_info)
        
    try:
        with open(session_path, 'w') as f:
            json.dump(session_data, f, indent=2)
    except Exception as e:
        print(f"Error writing session state for user {user_id}: {e}")

def log_agent_step(user_id: int, phase, message, logs_list, status="Success"):
    """Logs an agent step in memory, print to console, write to DB, and sync to file."""
    timestamp = datetime.now().strftime('%H:%M:%S')
    log_entry = {
        "timestamp": timestamp,
        "phase": phase,
        "message": message,
        "status": status
    }
    logs_list.append(log_entry)
    
    # Print to console
    print(f"[{timestamp}] {phase.upper()} (User {user_id}): {message}")
    
    # Save to SQLite/PostgreSQL AgentLog
    db = SessionLocal()
    try:
        db_log = AgentLog(
            user_id=user_id,
            action=phase,
            details=message[:500], # truncate to avoid huge DB records
            status=status
        )
        db.add(db_log)
        db.commit()
    except Exception as e:
        print(f"Failed to log step to DB: {e}")
    finally:
        db.close()

def replace_text_in_document(doc, replacements):
    """Finds and replaces placeholders in word documents, correctly applying formatting."""
    def add_formatted_text(paragraph, text):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                run = paragraph.add_run(part)

    for placeholder, content in replacements.items():
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder in para.text:
                            para.text = ""
                            for line in content.split('\n'):
                                add_formatted_text(para, line)
                                para.add_run('\n')

        for para in doc.paragraphs:
            if placeholder in para.text:
                para.text = ""
                for line in content.split('\n'):
                    add_formatted_text(para, line)
                    para.add_run('\n')

def ensure_tailored_documents(user_id, job_title, company_name, job_desc, logs_list, provider="gemini", model_name=None):
    """Ensures customized resume and cover letter exist in outputs. Creates them if missing."""
    safe_company = re.sub(r'[\\/*?:"<>|]', "", company_name)
    safe_job = re.sub(r'[\\/*?:"<>|]', "", job_title)
    
    final_dir = os.path.join(_PROJECT_ROOT, "outputs", str(user_id), safe_company, safe_job)
    os.makedirs(final_dir, exist_ok=True)
    
    resume_path = os.path.join(final_dir, "Resume.docx")
    cover_letter_path = os.path.join(final_dir, "CoverLetter.docx")
    
    if os.path.exists(resume_path) and os.path.exists(cover_letter_path):
        return resume_path, cover_letter_path
        
    log_agent_step(user_id, "thought", "Custom Resume/Cover Letter not found. Commencing real-time RAG tailoring...", logs_list)
    
    try:
        from core.profile_memory_resume_phase2 import generate_resume_and_reasoning, generate_coverletter_and_reasoning
        
        # 1. Tailor Resume
        log_agent_step(user_id, "action", f"Querying LLM ({provider}) to custom-tailor resume content via RAG...", logs_list)
        res_res = generate_resume_and_reasoning(job_title, job_desc, user_id, provider, model_name)
        doc_data = res_res.get("document", {})
        
        template_resume = os.path.join(_PROJECT_ROOT, "data", "Manu_Martin_Resume_Template1.docx")
        if os.path.exists(template_resume):
            doc = docx.Document(template_resume)
            replacements = {
                '{{SUMMARY}}': doc_data.get("summary", ""),
                '{{SKILLS}}': doc_data.get("skills", ""),
                '{{EXPERIENCE}}': doc_data.get("experience", ""),
                '{{PROJECTS}}': doc_data.get("projects", "")
            }
            replace_text_in_document(doc, replacements)
            doc.save(resume_path)
            log_agent_step(user_id, "observation", "Tailored resume compiled successfully and saved to outputs.", logs_list)
        else:
            log_agent_step(user_id, "observation", "Resume template not found at data/", logs_list, status="Failed")
            
        # 2. Tailor Cover Letter
        log_agent_step(user_id, "action", f"Querying LLM ({provider}) to custom-tailor cover letter...", logs_list)
        res_cl = generate_coverletter_and_reasoning(job_title, job_desc, user_id, provider, model_name)
        cl_text = res_cl.get("document", "")
        
        template_cl = os.path.join(_PROJECT_ROOT, "data", "Manu_Martin_coverletter_Template.docx")
        if os.path.exists(template_cl):
            doc_cl = docx.Document(template_cl)
            replace_text_in_document(doc_cl, {'{{COVER_LETTER_CONTENT}}': cl_text})
            doc_cl.save(cover_letter_path)
            log_agent_step(user_id, "observation", "Tailored cover letter compiled successfully and saved.", logs_list)
        else:
            log_agent_step(user_id, "observation", "Cover letter template not found.", logs_list, status="Failed")
            
    except Exception as err:
        log_agent_step(user_id, "observation", f"Document tailoring failed: {err}", logs_list, status="Failed")
        
    return resume_path, cover_letter_path

class AutonomousApplyAgent:
    def __init__(self, user_id: int, provider="gemini", model_name=None):
        self.user_id = user_id
        self.provider = provider
        self.model_name = model_name
        self.qdrant_client = get_qdrant_client()
        self.db = SessionLocal()

    def get_candidate_profile_context(self, job_title, job_desc):
        """Loads embedded facts from Qdrant and conversational facts from SQL database."""
        # 1. Qdrant RAG Context
        rag_context = get_rag_context(job_title, job_desc or "", self.qdrant_client, self.user_id)
        
        # 2. SQLite/PostgreSQL ProfileFacts for this user
        db_facts = self.db.query(ProfileFact).filter(ProfileFact.user_id == self.user_id).all()
        facts_str = ""
        if db_facts:
            facts_str = "## Stored Conversational Facts:\n" + "\n".join([f"- [{f.category}] {f.fact}" for f in db_facts])
            
        # 3. Base Fallback details if RAG is empty
        from core.db import UserProfile
        profile = self.db.query(UserProfile).filter_by(user_id=self.user_id).first()
        if profile:
            fallback_str = f"""
NAME: {profile.name or "Candidate"}
EMAIL: {profile.email or ""}
PHONE: {profile.phone or ""}
WEBSITE: {profile.website or ""}
GITHUB: {profile.github_username or ""}
"""
        else:
            fallback_str = """
NAME: Candidate
"""
        
        return f"{fallback_str}\n\n{rag_context}\n\n{facts_str}"

    def run_react_filling_cycle(self, job_url):
        # 1. Load job details
        job = self.db.query(Job).filter(Job.url == job_url, Job.user_id == self.user_id).first()
        job_title = job.title if job else "AI Engineer"
        company_name = job.company if job else "Company"
        job_desc = job.description if job else ""
        
        job_info = {
            "job_title": job_title,
            "company": company_name,
            "url": job_url
        }
        
        logs = []
        log_agent_step(self.user_id, "system", f"Autonomous Agent launched for '{job_title}' at '{company_name}'.", logs)
        update_session(self.user_id, "browsing", logs, job_info=job_info)
        
        profile_context = self.get_candidate_profile_context(job_title, job_desc)
        
        user_session_path = get_user_session_path(self.user_id)
        user_screenshot_path = get_user_screenshot_path(self.user_id)
        
        with sync_playwright() as p:
            log_agent_step(self.user_id, "action", "Launching Playwright browser context...", logs)
            
            # Headless mode is default so users see screenshot updates inline
            browser = p.chromium.launch(headless=True)
            context = browser.new_context(
                viewport={"width": 1280, "height": 800},
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            )
            page = context.new_page()
            stealth_sync(page)
            
            try:
                log_agent_step(self.user_id, "action", f"Navigating to job portal: {job_url}", logs)
                page.goto(job_url, timeout=30000)
                page.wait_for_load_state("networkidle")
                time.sleep(3) # safe padding
                
                # Check for cookies/dialogs and dismiss them
                try:
                    cookie_btn = page.locator("button:has-text('Accept'), button:has-text('Agree'), button:has-text('Cookies')").first
                    if cookie_btn.is_visible():
                        cookie_btn.click()
                        time.sleep(1)
                except:
                    pass
                
                # Handle potential "Apply Now" button click first
                apply_now = page.locator("button:has-text('Apply'), a:has-text('Apply Now'), button:has-text('Apply Now')").first
                if apply_now.is_visible() and not page.locator("input[name*='name'], input[type='email']").first.is_visible():
                    log_agent_step(self.user_id, "action", "Found 'Apply Now' gateway, clicking to open application form...", logs)
                    apply_now.click()
                    time.sleep(3)
                    page.wait_for_load_state("networkidle")
                
                # Dynamic page processing loop (up to 3 pages)
                for page_idx in range(1, 4):
                    log_agent_step(self.user_id, "thought", f"Scanning page structure (Form Page {page_idx})...", logs)
                    page.screenshot(path=user_screenshot_path)
                    update_session(self.user_id, "browsing", logs, user_screenshot_path)
                    
                    # Scrape DOM elements
                    fields = page.evaluate("""() => {
                        const items = [];
                        document.querySelectorAll("input, textarea, select, button").forEach(el => {
                            // Skip hidden elements, submits, and standard navigation links
                            if (el.tagName === 'INPUT' && ['hidden', 'submit', 'image'].includes(el.type)) return;
                            if (el.tagName === 'BUTTON' && ['submit'].includes(el.type)) return;
                            
                            // Check visibility
                            const rect = el.getBoundingClientRect();
                            if (rect.width === 0 || rect.height === 0) return;
                            
                            // Retrieve label text
                            let labelText = "";
                            if (el.id) {
                                const label = document.querySelector(`label[for="${el.id}"]`);
                                if (label) labelText = label.innerText;
                            }
                            if (!labelText) {
                                const label = el.closest('label');
                                if (label) {
                                    labelText = label.innerText;
                                } else {
                                    const parent = el.parentElement;
                                    if (parent) {
                                        labelText = parent.innerText.split('\\n')[0];
                                    }
                                }
                            }
                            
                            // Select options
                            let options = [];
                            if (el.tagName === 'SELECT') {
                                options = Array.from(el.options).map(o => ({ value: o.value, text: o.text }));
                            }
                            
                            items.push({
                                tagName: el.tagName,
                                type: el.type || '',
                                name: el.name || '',
                                id: el.id || '',
                                placeholder: el.placeholder || '',
                                labelText: labelText.trim(),
                                required: el.required || false,
                                options: options,
                                text: el.innerText || ''
                            });
                        });
                        return items;
                    }""")
                    
                    if not fields:
                        log_agent_step(self.user_id, "observation", "No form elements detected on the current view.", logs)
                        break
                        
                    log_agent_step(self.user_id, "observation", f"Detected {len(fields)} interactive form elements.", logs)
                    
                    # Prompt LLM for form filling decisions
                    log_agent_step(self.user_id, "thought", "Querying LLM to formulate form-filling commands...", logs)
                    update_session(self.user_id, "reasoning", logs, user_screenshot_path)
                    
                    llm = get_llm(provider=self.provider, model_name=self.model_name)
                    
                    prompt = f"""
You are Eva, the autonomous browser solver agent. Your task is to analyze the form fields on the current webpage and determine how to fill them using the candidate's profile facts.

CANDIDATE PROFILE CONTEXT:
---
{profile_context}
---

JOB DETAILS:
Job Title: {job_title}
Company: {company_name}
Job Description:
---
{job_desc}
---

CURRENT FORM FIELDS:
---
{json.dumps(fields, indent=2)}
---

INSTRUCTIONS:
1. Match each form field (by its list index) to the candidate's profile facts.
2. For text inputs/textareas, write a professional, tailored response if the field asks for custom answers (e.g. why do you want to work here, why are you a good fit, etc.). For standard fields (name, email, phone, linkedin, github, portfolio, location), use the exact candidate details.
3. For file inputs:
   - If the label or name mentions "resume" or "cv", output action "upload_resume".
   - If the label or name mentions "cover letter", output action "upload_cover_letter".
4. For checkboxes, select dropdowns, or radio buttons, output "click" or "select" action and decide the appropriate value/state.
5. Identify any "Next", "Continue", or "Submit" buttons in the field list.
   - If it is a "Next" or "Continue" button, you may decide to click it AFTER filling out the page elements.
   - If it is a "Submit" button, do NOT click it. Instead, we will wait for human approval.
6. Return a VALID JSON object with the following schema:
{{
  "thought": "Brief overall thought about the fields on this page.",
  "actions": [
    {{
      "index": 0, // integer index of the element in the list
      "action": "fill | click | select | upload_resume | upload_cover_letter",
      "value": "Value to fill/select, or empty for click/upload",
      "thought": "Reason for doing this."
    }}
  ]
}}
Do NOT output any markdown tags or text outside of the JSON object.
"""
                    response = llm.invoke(prompt)
                    content = response.content if hasattr(response, 'content') else str(response)
                    parsed_json = parse_json_output(content)
                    
                    if not parsed_json or "actions" not in parsed_json:
                        log_agent_step(self.user_id, "observation", "Failed to parse form instructions from LLM. Aborting.", logs, status="Failed")
                        update_session(self.user_id, "failed", logs, user_screenshot_path)
                        return
                        
                    log_agent_step(self.user_id, "thought", f"AI Decision: {parsed_json.get('thought')}", logs)
                    
                    actions = parsed_json["actions"]
                    
                    # Execute filling actions
                    has_navigated = False
                    for act in actions:
                        idx = act["index"]
                        action_type = act["action"]
                        val = act.get("value", "")
                        thought = act.get("thought", "")
                        
                        if idx < 0 or idx >= len(fields):
                            continue
                            
                        field = fields[idx]
                        field_name = field['labelText'] or field['name'] or field['placeholder'] or f"Element {idx}"
                        
                        log_agent_step(self.user_id, "thought", f"Action planning for '{field_name}': {thought}", logs)
                        
                        try:
                            locator = page.locator("input, textarea, select, button").nth(idx)
                            
                            if action_type == "fill":
                                locator.fill(val)
                                log_agent_step(self.user_id, "action", f"Filled '{field_name}' with '{val}'", logs)
                            elif action_type == "select":
                                locator.select_option(val)
                                log_agent_step(self.user_id, "action", f"Selected '{val}' for '{field_name}'", logs)
                            elif action_type == "upload_resume":
                                resume_path, _ = ensure_tailored_documents(
                                    self.user_id, job_title, company_name, job_desc, logs, self.provider, self.model_name
                                )
                                if resume_path and os.path.exists(resume_path):
                                    locator.set_input_files(resume_path)
                                    log_agent_step(self.user_id, "action", f"Uploaded tailored resume: {os.path.basename(resume_path)}", logs)
                                else:
                                    log_agent_step(self.user_id, "observation", "Tailored resume upload failed.", logs, status="Failed")
                            elif action_type == "upload_cover_letter":
                                _, cl_path = ensure_tailored_documents(
                                    self.user_id, job_title, company_name, job_desc, logs, self.provider, self.model_name
                                )
                                if cl_path and os.path.exists(cl_path):
                                    locator.set_input_files(cl_path)
                                    log_agent_step(self.user_id, "action", f"Uploaded tailored cover letter: {os.path.basename(cl_path)}", logs)
                                else:
                                    log_agent_step(self.user_id, "observation", "Tailored cover letter upload failed.", logs, status="Failed")
                            elif action_type == "click":
                                # If clicking next/continue, make sure it happens at the end
                                if any(kw in field_name.lower() for kw in ["next", "continue", "submit"]):
                                    has_navigated = True
                                locator.click()
                                log_agent_step(self.user_id, "action", f"Clicked: '{field_name}'", logs)
                                time.sleep(2)
                                
                            page.screenshot(path=user_screenshot_path)
                            update_session(self.user_id, "filling", logs, user_screenshot_path)
                            time.sleep(0.5)
                            
                        except Exception as field_err:
                            log_agent_step(self.user_id, "observation", f"Error filling field '{field_name}': {field_err}", logs, status="Failed")
                    
                    if not has_navigated:
                        # If no navigation clicked, we completed the form details on the page
                        break
                        
                # End of filling loop, now wait for Human-in-the-Loop review
                log_agent_step(self.user_id, "observation", "All form fields filled. Transitioning to Human review mode.", logs)
                page.screenshot(path=user_screenshot_path)
                update_session(self.user_id, "waiting_approval", logs, user_screenshot_path, command="")
                
                # Poll outputs/{user_id}/agent_session.json for approval command
                timeout_seconds = 300 # 5 minutes
                start_time = time.time()
                approved = False
                
                while time.time() - start_time < timeout_seconds:
                    if page.is_closed():
                        log_agent_step(self.user_id, "observation", "Browser viewport closed by user.", logs, status="Failed")
                        break
                        
                    # Read command from file
                    try:
                        with open(user_session_path, 'r') as f:
                            s_data = json.load(f)
                            cmd = s_data.get("command")
                            if cmd == "approve":
                                approved = True
                                break
                            elif cmd == "abort":
                                log_agent_step(self.user_id, "action", "User sent abort signal. Shutting down browser session.", logs, status="Failed")
                                break
                    except:
                        pass
                        
                    time.sleep(1.5)
                
                if approved:
                    log_agent_step(self.user_id, "action", "Submission approved! Locating and clicking final Submit button...", logs)
                    update_session(self.user_id, "submitting", logs, user_screenshot_path)
                    
                    try:
                        # Locate submit button
                        submit_btn = page.locator("button[type='submit'], input[type='submit'], button:has-text('Submit'), button:has-text('Submit Application')").first
                        if submit_btn.is_visible():
                            submit_btn.click()
                            time.sleep(5) # Wait for processing
                            page.wait_for_load_state("networkidle")
                            log_agent_step(self.user_id, "observation", "Application submitted successfully!", logs)
                            
                            # Mark job in DB as applied
                            if job:
                                job.is_applied = True
                                # Create Application record
                                app_record = self.db.query(Application).filter_by(job_id=job.id).first()
                                if not app_record:
                                    app_record = Application(job_id=job.id, status="Applied")
                                    self.db.add(app_record)
                                else:
                                    app_record.status = "Applied"
                                self.db.commit()
                                
                            page.screenshot(path=user_screenshot_path)
                            update_session(self.user_id, "success", logs, user_screenshot_path)
                        else:
                            log_agent_step(self.user_id, "observation", "Could not locate final Submit button on page.", logs, status="Failed")
                            update_session(self.user_id, "failed", logs, user_screenshot_path)
                    except Exception as submit_err:
                        log_agent_step(self.user_id, "observation", f"Error during click submission: {submit_err}", logs, status="Failed")
                        update_session(self.user_id, "failed", logs, user_screenshot_path)
                else:
                    update_session(self.user_id, "failed", logs, user_screenshot_path)
                    
            except Exception as e:
                log_agent_step(self.user_id, "observation", f"Agent execution encountered critical error: {e}", logs, status="Failed")
                update_session(self.user_id, "failed", logs, user_screenshot_path)
            finally:
                browser.close()
                self.db.close()

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Autonomous Playwright Solver Agent")
    parser.add_argument("--url", type=str, required=True, help="Job URL to apply to")
    parser.add_argument("--user-id", type=int, required=True, help="ID of the user running this task")
    parser.add_argument("--provider", type=str, default="gemini", help="LLM Provider")
    parser.add_argument("--model", type=str, default=None, help="LLM Model name")
    args = parser.parse_args()

    agent = AutonomousApplyAgent(user_id=args.user_id, provider=args.provider, model_name=args.model)
    agent.run_react_filling_cycle(args.url)
