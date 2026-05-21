# dashboard/coverletter_resume_app_phase4.py
'''
run in cmd: 
"C:/Program Files (x86)/Microsoft/Edge/Application/msedge.exe" -remote-debugging-port=9222
'''
import streamlit as st
import pandas as pd
import os
import re
import docx
from docx2pdf import convert
import pythoncom
from st_aggrid import AgGrid, GridOptionsBuilder
import subprocess
import sys

# --- Setup Paths ---
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(_SCRIPT_DIR)
OUTPUTS_DIR = os.path.join(_PROJECT_ROOT, "outputs")
JOBS_DIR = os.path.join(_PROJECT_ROOT, "data", "jobs_details")
DATA_DIR = os.path.join(_PROJECT_ROOT, "data")

# --- Import your AI functions ---
from core.profile_memory_resume_phase2 import generate_resume_and_reasoning, generate_coverletter_and_reasoning, ingest_and_embed

# --- Robust Parsing and Template Functions ---
def get_latest_csv_file(path):
    """Finds the most recently created CSV file in a directory."""
    if not os.path.exists(path):
        return None, []
    
    csv_files = [f for f in os.listdir(path) if f.endswith('.csv')]
    if not csv_files:
        return None, []
        
    # Get creation times and find the latest file
    full_paths = [os.path.join(path, f) for f in csv_files]
    latest_file = max(full_paths, key=os.path.getctime)
    
    return latest_file, sorted(csv_files,reverse=True)

def replace_text_in_document(doc, placeholder, content):
    """
    Finds and replaces a placeholder, correctly applying bold formatting
    for any text surrounded by double asterisks (e.g., **text**).
    """
    # This helper function does the complex work of adding formatted text
    def add_formatted_text(paragraph, text):
        # Split the text by the bold markers, keeping the markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # If the part is marked for bold, add it as a bold run
                # and remove the asterisks
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Otherwise, add it as a normal run
                run = paragraph.add_run(part)

    # --- Main logic to search the whole document ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        # Clear the placeholder text
                        para.text = ""
                        # Add the new, formatted content line by line
                        for line in content.split('\n'):
                            add_formatted_text(para, line)
                            para.add_run('\n') # Add line breaks back

    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = ""
            for line in content.split('\n'):
                add_formatted_text(para, line)
                para.add_run('\n')



def parse_ai_response_for_template_stars(text: str) -> dict:
    """Parses the AI's markdown-style output into a dictionary for the template."""
    sections = {}
    headers = ["Professional Summary", "Key Skills", "Relevant Projects", "Key Projects", "Relevant Experience", "Experience"]
    pattern = r"\*\*((" + "|".join(headers) + r"))\*\*"
    parts = re.split(pattern, text)
    for i in range(1, len(parts), 3):
        header = parts[i].strip()
        content = parts[i+2].strip()
        if "summary" in header.lower(): sections['{{SUMMARY}}'] = content
        elif "skills" in header.lower(): sections['{{SKILLS}}'] = content
        elif "projects" in header.lower(): sections['{{PROJECTS}}'] = content
        elif "experience" in header.lower(): sections['{{EXPERIENCE}}'] = content
    return sections

def parse_ai_response_for_template_hastag(text: str) -> dict:
    """Parses the AI's markdown-style output into a dictionary for the template."""
    sections = {}
    headers = ["Professional Summary", "Key Skills", "Relevant Projects", "Key Projects", "Relevant Experience", "Experience"]
    pattern = r"\#\#\# ((" + "|".join(headers) + r"))"
    # print(f"pattern: {pattern}")
    parts = re.split(pattern, text)
    # print(f"part: {parts}")
    # print(f"len(parts): {len(parts)}")
    for i in range(1, len(parts), 3):
        header = parts[i].strip()
        content = parts[i+2].strip()
        if "summary" in header.lower(): sections['{{SUMMARY}}'] = content
        elif "skills" in header.lower(): sections['{{SKILLS}}'] = content
        elif "projects" in header.lower(): sections['{{PROJECTS}}'] = content
        elif "experience" in header.lower(): sections['{{EXPERIENCE}}'] = content
    return sections

def parse_ai_response_for_coverletter_template(text: str) -> dict:
    """A robust parser for the AI's markdown-style output."""
    sections = {}
    sections['{{COVER_LETTER_CONTENT}}'] = text
    return sections

# --- Streamlit App ---
st.set_page_config(layout="wide")
st.title("🤖 Eva - AI Job Butler")

if 'selected_job_data' not in st.session_state:
    st.session_state.selected_job_data = ""
if 'job_title' not in st.session_state:
    st.session_state.job_title = ""
if 'company_name' not in st.session_state:
    st.session_state.company_name = ""
if 'job_desc' not in st.session_state:
    st.session_state.job_desc = ""
if 'job_url' not in st.session_state:
    st.session_state.job_url = ""

# --- Initialize Session State ---
if 'document_draft_resume' not in st.session_state:
    st.session_state.document_draft_resume = ""
    st.session_state.reasoning_resume = ""
if 'document_draft_coverletter' not in st.session_state:
    st.session_state.document_draft_coverletter = ""
    st.session_state.reasoning_coverletter = ""

first_col1, first_col2 = st.columns([3,1])
with first_col1:

    st.header("1. Job Selection")

    latest_csv, all_csvs = get_latest_csv_file(JOBS_DIR)

    if not latest_csv:
        st.warning(f"No job CSVs found in '{JOBS_DIR}'. Please run your scraper first.")
    else:
        # Allow user to select a CSV file, defaulting to the latest one
        selected_csv_name = st.selectbox("Select Job Lead File", options=all_csvs, index=all_csvs.index(os.path.basename(latest_csv)))
        jobs_df = pd.read_csv(os.path.join(JOBS_DIR, selected_csv_name))                # Load the selected CSV
        # st.info("Click on a row to select a job and populate the details below.")
        jobs_df_processed = jobs_df[['job_url','job_url_direct','title','company','location','date_posted','description']].copy()
        # --- AgGrid Implementation ---
        gb = GridOptionsBuilder.from_dataframe(jobs_df_processed)
        gb.configure_selection(
            'single',  # Single row selection
            use_checkbox=False,  # <-- Key fix: No checkboxes for single mode (enables proper row-click selection)
            rowMultiSelectWithClick=False, 
            suppressRowDeselection=False
        )
        gb.configure_grid_options(domLayout='normal')
        gridOptions = gb.build()

        grid_response = AgGrid(
            jobs_df_processed,
            gridOptions=gridOptions,
            height=400,
            width='100%',
            allow_unsafe_jscode=True,
            # enable_enterprise_modules=False,
            # update_on=['selection_changed'],
            fit_columns_on_grid_load=True,
            theme='streamlit'
        )

        # Capture the selected row
        selected_rows = grid_response['selected_rows']
        selected = grid_response['selected_rows']
        st.write("Selected:", selected)

        print(f"selected_rows: {selected_rows}")
        if selected_rows is not None and not selected_rows.empty:
            # When a row is clicked, store its data in session state
            selected_job_data = selected_rows.to_dict('records')[0]
            st.session_state.selected_job_data = selected_job_data  # Store full row for consistency
            st.session_state.job_title = selected_job_data.get('title', '')
            st.session_state.company_name = selected_job_data.get('company', '')
            st.session_state.job_desc = selected_job_data.get('description', '')
            st.session_state.job_url = selected_job_data.get("job_url","")
            
            # print(f"Updated session state - Job Title: {st.session_state.job_title}")  # Optional: Log for debugging
            # No st.rerun() needed - AgGrid handles it!

with first_col2:
    st.info("Please update and configure the necessary details in the core\ingest.py file and then tap below 👇")
    if st.button("Ingest Profile"):
        ingest_and_embed(clear_existing=True)  # Initial ingest
    
        st.success("Profile ingested and embedded!")

    st.markdown("---")

    search_term = st.text_input("Job Title to Search", "AI Engineer")
    location = st.text_input("Location", "Bengaluru")

    if st.button("Fetch Latest Jobs"):
        # --- THIS IS THE NEW, CORRECTED LOGIC ---
        st.info("Starting the job scout... This may take a moment.")
        
        # We will display the live output from the scraper script
        with st.expander("Show Scraper Log"):
            # Use subprocess.Popen to run the script as a separate process
            process = subprocess.Popen(
                [sys.executable, "scrapers/job_scout.py", "--search", search_term, "--location", location],
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                encoding='utf-8'
            )
            
            # Display the output in real-time
            log_container = st.empty()
            log_text = ""
            while True:
                output = process.stdout.readline()
                if output == '' and process.poll() is not None:
                    break
                if output:
                    log_text += output
                    log_container.code(log_text)
        
        st.success("Job search process complete! The latest jobs have been saved to the database.")



# --- Main Layout ---
col1, col2, col3 = st.columns([1, 1, 1])

# --- Column 1: Inputs ---
with col1:

    job_data = st.session_state.selected_job_data if st.session_state.selected_job_data else {}

    # selected_job = st.session_state.selected_job
    st.header("1. Job Details")
    if st.session_state.selected_job_data and isinstance(st.session_state.selected_job_data, dict):    
        if st.session_state.selected_job_data.get('job_url'):
            st.markdown(f"Job Link: [Link]({st.session_state.selected_job_data.get('job_url')})")
        if st.session_state.selected_job_data.get('job_url_direct'):
            st.markdown(f"Job direct Link: [link]({st.session_state.selected_job_data.get('job_url_direct')})")
        job_title = st.text_input("Job Title", key='job_title')
    company_name = st.text_input("Company Name", key="company_name")
    job_desc_input = st.text_area("Paste Full Job Description", key='job_desc', height=250)

    need_resume = st.checkbox("Generate Resume", value=True)
    need_coverletter = st.checkbox("Generate Cover Letter", value=True)


    if st.button("✨ Generate Drafts", use_container_width=True):
        if job_desc_input and job_title:
            with st.spinner("Eva is reasoning and crafting all documents..."):
                # Generate Resume
                if need_resume:
                    resume_dict = generate_resume_and_reasoning(job_title, job_desc_input)
                    st.session_state.document_draft_resume = resume_dict.get('document', '')
                    st.session_state.reasoning_resume = resume_dict.get('reasoning', '')
                # Generate Cover Letter
                if need_coverletter:
                    cover_dict = generate_coverletter_and_reasoning(job_title, job_desc_input)
                    st.session_state.document_draft_coverletter = cover_dict.get('document', '')
                    st.session_state.reasoning_coverletter = cover_dict.get('reasoning', '')
        else:
            st.warning("Please provide a Job Title and Description.")


# --- Column 2: Resume Editor ---
with col2:
    st.header("2. Review Resume")
    if st.session_state.document_draft_resume:
        resume_edited_text = st.text_area("Resume Editor", st.session_state.document_draft_resume, height=400, key="resume_editor")
        with st.expander("Show Resume Reasoning 💡"):
            st.info(st.session_state.reasoning_resume)

        if st.button("✅ Approve & Generate Resume PDF", use_container_width=True):
            try:
                template_path = os.path.join(DATA_DIR, "Manu_Martin_Resume_Template1.docx")
                # ... (your folder creation logic) ...
                company_folder = re.sub(r'[\\/*?:"<>|]', "", company_name)
                job_folder = re.sub(r'[\\/*?:"<>|]', "", job_title)
                final_dir = os.path.join(OUTPUTS_DIR, company_folder, job_folder)
                os.makedirs(final_dir, exist_ok=True)
                final_pdf_path = os.path.join(final_dir, "final_resume.pdf") # Simplified for example
                temp_docx_path = os.path.join(final_dir, "temp_resume.docx")

                doc = docx.Document(template_path)
                sections_to_fill = parse_ai_response_for_template_hastag(resume_edited_text)
                # print(f"edited_text: {edited_text}")                  # For debugging
                # print(f"Sections to fill: {sections_to_fill}")        # For debugging
                if not sections_to_fill:
                    sections_to_fill = parse_ai_response_for_template_stars(resume_edited_text)
                    print(f"Sections to fill after stars parsing: {sections_to_fill}")

                print(f"sections to fill:{sections_to_fill}, resume_edited test:{resume_edited_text} ")

                for placeholder, content in sections_to_fill.items():
                    replace_text_in_document(doc, placeholder, content)
                
                doc.save(temp_docx_path)
                
                pythoncom.CoInitialize()
                convert(temp_docx_path, final_pdf_path)
                os.remove(temp_docx_path)

                st.success(f"Resume PDF saved!")
                with open(final_pdf_path, "rb") as pdf_file:
                    st.download_button("Download Resume PDF", pdf_file.read(), file_name="final_resume.pdf")
            except Exception as e:
                st.error(f"Failed to generate resume PDF: {e}")
            finally:
                pythoncom.CoUninitialize()

# --- Column 3: Cover Letter Editor ---
with col3:
    st.header("3. Review Cover Letter")
    if st.session_state.document_draft_coverletter:
        coverletter_edited_text = st.text_area("Cover Letter Editor", st.session_state.document_draft_coverletter, height=400, key="coverletter_editor")
        with st.expander("Show Cover Letter Reasoning 💡"):
            st.info(st.session_state.reasoning_coverletter)
        
        if st.button("✅ Approve & Generate Coverletter PDF", use_container_width=True):
            try:
                template_path = os.path.join(DATA_DIR, "Manu_Martin_Coverletter_Template.docx")
                # ... (your folder creation logic) ...
                company_folder = re.sub(r'[\\/*?:"<>|]', "", company_name)
                job_folder = re.sub(r'[\\/*?:"<>|]', "", job_title)
                final_dir = os.path.join(OUTPUTS_DIR, company_folder, job_folder)
                os.makedirs(final_dir, exist_ok=True)
                final_pdf_path = os.path.join(final_dir, "final_coverletter.pdf") # Simplified for example
                temp_docx_path = os.path.join(final_dir, "temp_coverletter.docx")

                doc = docx.Document(template_path)
                sections_to_fill = parse_ai_response_for_coverletter_template(coverletter_edited_text) # Use the edited text
                print(f"section to fill: {sections_to_fill}")
                print(f"coverletter Edited Text: {coverletter_edited_text}")

                for placeholder, content in sections_to_fill.items():
                    replace_text_in_document(doc, placeholder, content)
                
                doc.save(temp_docx_path)
                
                pythoncom.CoInitialize()
                convert(temp_docx_path, final_pdf_path)
                os.remove(temp_docx_path)

                st.success(f"Cover Letter PDF saved!")
                with open(final_pdf_path, "rb") as pdf_file:
                    st.download_button("Download Cover Letter PDF", pdf_file.read(), file_name="final_coverletter.pdf")
            except Exception as e:
                st.error(f"Failed to generate coverletter PDF: {e}")
            finally:
                pythoncom.CoUninitialize()

        # Add a similar "Approve & Generate PDF" button for the cover letter here if needed