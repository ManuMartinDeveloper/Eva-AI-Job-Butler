# dashboard/coverletter_resume_app_phase4.py
import streamlit as st
import os
import re
import docx
from docx2pdf import convert
import pythoncom

# --- Setup Paths ---
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(_SCRIPT_DIR)
OUTPUTS_DIR = os.path.join(_PROJECT_ROOT, "outputs")
DATA_DIR = os.path.join(_PROJECT_ROOT, "data")

# --- Import your AI functions ---
from core.profile_memory_resume_phase2 import generate_resume_and_reasoning, generate_coverletter_and_reasoning

# --- Robust Parsing and Template Functions ---

def replace_text_in_document(doc, placeholder, content):
    """
    Finds and replaces a placeholder, correctly applying bold formatting
    for any text surrounded by double asterisks (e.g., **text**).
    """
    # This helper function does the complex work of adding formatted text
    def add_formatted_text(paragraph, text):
        # Split the text by the bold markers, keeping the markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # If the part is marked for bold, add it as a bold run
                # and remove the asterisks
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Otherwise, add it as a normal run
                run = paragraph.add_run(part)

    # --- Main logic to search the whole document ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        # Clear the placeholder text
                        para.text = ""
                        # Add the new, formatted content line by line
                        for line in content.split('\n'):
                            add_formatted_text(para, line)
                            para.add_run('\n') # Add line breaks back

    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = ""
            for line in content.split('\n'):
                add_formatted_text(para, line)
                para.add_run('\n')



def parse_ai_response_for_template_stars(text: str) -> dict:
    """Parses the AI's markdown-style output into a dictionary for the template."""
    sections = {}
    headers = ["Professional Summary", "Key Skills", "Relevant Projects", "Key Projects", "Relevant Experience", "Experience"]
    pattern = r"\*\*((" + "|".join(headers) + r"))\*\*"
    parts = re.split(pattern, text)
    for i in range(1, len(parts), 3):
        header = parts[i].strip()
        content = parts[i+2].strip()
        if "summary" in header.lower(): sections['{{SUMMARY}}'] = content
        elif "skills" in header.lower(): sections['{{SKILLS}}'] = content
        elif "projects" in header.lower(): sections['{{PROJECTS}}'] = content
        elif "experience" in header.lower(): sections['{{EXPERIENCE}}'] = content
    return sections

def parse_ai_response_for_template_hastag(text: str) -> dict:
    """Parses the AI's markdown-style output into a dictionary for the template."""
    sections = {}
    headers = ["Professional Summary", "Key Skills", "Relevant Projects", "Key Projects", "Relevant Experience", "Experience"]
    pattern = r"\#\#\# ((" + "|".join(headers) + r"))"
    # print(f"pattern: {pattern}")
    parts = re.split(pattern, text)
    # print(f"part: {parts}")
    # print(f"len(parts): {len(parts)}")
    for i in range(1, len(parts), 3):
        header = parts[i].strip()
        content = parts[i+2].strip()
        if "summary" in header.lower(): sections['{{SUMMARY}}'] = content
        elif "skills" in header.lower(): sections['{{SKILLS}}'] = content
        elif "projects" in header.lower(): sections['{{PROJECTS}}'] = content
        elif "experience" in header.lower(): sections['{{EXPERIENCE}}'] = content
    return sections

def parse_ai_response_for_coverletter_template(text: str) -> dict:
    """A robust parser for the AI's markdown-style output."""
    sections = {}
    sections['{{COVER_LETTER_CONTENT}}'] = text
    return sections

# --- Streamlit App ---
st.set_page_config(layout="wide")
st.title("🤖 Eva - AI Job Butler")

# --- Initialize Session State ---
if 'document_draft_resume' not in st.session_state:
    st.session_state.document_draft_resume = ""
    st.session_state.reasoning_resume = ""
if 'document_draft_coverletter' not in st.session_state:
    st.session_state.document_draft_coverletter = ""
    st.session_state.reasoning_coverletter = ""

# --- Main Layout ---
col1, col2, col3 = st.columns([1, 1.5, 1.5])

# --- Column 1: Inputs ---
with col1:
    st.header("1. Job Details")
    job_title = st.text_input("Job Title", "AI Engineer")
    company_name = st.text_input("Company Name", "Tech Innovations Inc.")
    job_desc_input = st.text_area("Paste Full Job Description", height=250)

    if st.button("✨ Generate Drafts", use_container_width=True):
        if job_desc_input and job_title:
            with st.spinner("Eva is reasoning and crafting all documents..."):
                # Generate Resume
                resume_dict = generate_resume_and_reasoning(job_title, job_desc_input)
                st.session_state.document_draft_resume = resume_dict.get('document', '')
                st.session_state.reasoning_resume = resume_dict.get('reasoning', '')
                # Generate Cover Letter
                cover_dict = generate_coverletter_and_reasoning(job_title, job_desc_input)
                st.session_state.document_draft_coverletter = cover_dict.get('document', '')
                st.session_state.reasoning_coverletter = cover_dict.get('reasoning', '')
        else:
            st.warning("Please provide a Job Title and Description.")

# --- Column 2: Resume Editor ---
with col2:
    st.header("2. Review Resume")
    if st.session_state.document_draft_resume:
        resume_edited_text = st.text_area("Resume Editor", st.session_state.document_draft_resume, height=400, key="resume_editor")
        with st.expander("Show Resume Reasoning 💡"):
            st.info(st.session_state.reasoning_resume)

        if st.button("✅ Approve & Generate Resume PDF", use_container_width=True):
            try:
                template_path = os.path.join(DATA_DIR, "Manu_Martin_Resume_Template1.docx")
                # ... (your folder creation logic) ...
                company_folder = re.sub(r'[\\/*?:"<>|]', "", company_name)
                job_folder = re.sub(r'[\\/*?:"<>|]', "", job_title)
                final_dir = os.path.join(OUTPUTS_DIR, company_folder, job_folder)
                os.makedirs(final_dir, exist_ok=True)
                final_pdf_path = os.path.join(OUTPUTS_DIR, "final_resume.pdf") # Simplified for example
                temp_docx_path = os.path.join(OUTPUTS_DIR, "temp_resume.docx")

                doc = docx.Document(template_path)
                sections_to_fill = parse_ai_response_for_template_hastag(resume_edited_text)
                # print(f"edited_text: {edited_text}")                  # For debugging
                # print(f"Sections to fill: {sections_to_fill}")        # For debugging
                if not sections_to_fill:
                    sections_to_fill = parse_ai_response_for_template_stars(resume_edited_text)
                    print(f"Sections to fill after stars parsing: {sections_to_fill}")

                print(f"sections to fill:{sections_to_fill}, resume_edited test:{resume_edited_text} ")

                for placeholder, content in sections_to_fill.items():
                    replace_text_in_document(doc, placeholder, content)
                
                doc.save(temp_docx_path)
                
                pythoncom.CoInitialize()
                convert(temp_docx_path, final_pdf_path)
                os.remove(temp_docx_path)

                st.success(f"Resume PDF saved!")
                with open(final_pdf_path, "rb") as pdf_file:
                    st.download_button("Download Resume PDF", pdf_file.read(), file_name="final_resume.pdf")
            except Exception as e:
                st.error(f"Failed to generate resume PDF: {e}")
            finally:
                pythoncom.CoUninitialize()

# --- Column 3: Cover Letter Editor ---
with col3:
    st.header("3. Review Cover Letter")
    if st.session_state.document_draft_coverletter:
        coverletter_edited_text = st.text_area("Cover Letter Editor", st.session_state.document_draft_coverletter, height=400, key="coverletter_editor")
        with st.expander("Show Cover Letter Reasoning 💡"):
            st.info(st.session_state.reasoning_coverletter)
        
        if st.button("✅ Approve & Generate Coverletter PDF", use_container_width=True):
            try:
                template_path = os.path.join(DATA_DIR, "Manu_Martin_Coverletter_Template.docx")
                # ... (your folder creation logic) ...
                final_pdf_path = os.path.join(OUTPUTS_DIR, "final_coverletter.pdf") # Simplified for example
                temp_docx_path = os.path.join(OUTPUTS_DIR, "temp_coverletter.docx")

                doc = docx.Document(template_path)
                sections_to_fill = parse_ai_response_for_coverletter_template(coverletter_edited_text) # Use the edited text
                print(f"section to fill: {sections_to_fill}")
                print(f"coverletter Edited Text: {coverletter_edited_text}")

                for placeholder, content in sections_to_fill.items():
                    replace_text_in_document(doc, placeholder, content)
                
                doc.save(temp_docx_path)
                
                pythoncom.CoInitialize()
                convert(temp_docx_path, final_pdf_path)
                os.remove(temp_docx_path)

                st.success(f"Cover Letter PDF saved!")
                with open(final_pdf_path, "rb") as pdf_file:
                    st.download_button("Download Cover Letter PDF", pdf_file.read(), file_name="final_coverletter.pdf")
            except Exception as e:
                st.error(f"Failed to generate coverletter PDF: {e}")
            finally:
                pythoncom.CoUninitialize()

        # Add a similar "Approve & Generate PDF" button for the cover letter here if needed