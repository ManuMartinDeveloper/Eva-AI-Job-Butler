import streamlit as st
import os
import sys
import re
import docx
import json

# --- Setup Paths ---
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(os.path.dirname(_SCRIPT_DIR))
sys.path.append(_PROJECT_ROOT)

from core.profile_memory_resume_phase2 import generate_resume_and_reasoning, generate_coverletter_and_reasoning

OUTPUTS_DIR = os.path.join(_PROJECT_ROOT, "outputs")
DATA_DIR = os.path.join(_PROJECT_ROOT, "data")

st.set_page_config(page_title="Document Generator - Eva", page_icon="✍️", layout="wide")

st.title("✍️ Document Generator")

# --- Helper Functions ---
def replace_text_in_document(doc, placeholder, content):
    """
    Finds and replaces a placeholder, correctly applying bold formatting.
    """
    def add_formatted_text(paragraph, text):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                run = paragraph.add_run(part)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        para.text = ""
                        for line in content.split('\n'):
                            add_formatted_text(para, line)
                            para.add_run('\n')

    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = ""
            for line in content.split('\n'):
                add_formatted_text(para, line)
                para.add_run('\n')

# --- Session State ---
if 'selected_job' not in st.session_state:
    st.session_state.selected_job = {}

if 'generated_resume' not in st.session_state:
    st.session_state.generated_resume = {}
if 'generated_coverletter' not in st.session_state:
    st.session_state.generated_coverletter = {}

# --- Main UI ---
col1, col2 = st.columns([1, 1])

with col1:
    st.header("1. Job Details")
    
    selected_job = st.session_state.selected_job
    
    if selected_job:
        st.success(f"Selected Job: {selected_job.get('title')} @ {selected_job.get('company')}")
        job_title = st.text_input("Job Title", value=selected_job.get('title', ''))
        company_name = st.text_input("Company Name", value=selected_job.get('company', ''))
        # job_desc = st.text_area("Job Description", value=selected_job.get('description', ''), height=200)
        # Use a placeholder if description is missing
        desc_val = selected_job.get('description', '')
        if isinstance(desc_val, float): # Handle NaN
             desc_val = ""
        job_desc = st.text_area("Job Description", value=desc_val, height=200)
    else:
        st.warning("No job selected. Go to 'Home' to select a job or enter details manually.")
        job_title = st.text_input("Job Title")
        company_name = st.text_input("Company Name")
        job_desc = st.text_area("Job Description", height=200)

    # Settings retrieval
    provider = st.session_state.get('llm_provider', 'gemini')
    model = st.session_state.get('llm_model', None)

    if st.button("✨ Generate Documents", use_container_width=True):
        if job_title and job_desc:
            with st.spinner(f"Eva is thinking... (Using {provider})"):
                # Generate Resume
                resume_result = generate_resume_and_reasoning(job_title, job_desc, provider=provider, model_name=model)
                st.session_state.generated_resume = resume_result
                
                # Generate Cover Letter
                cover_result = generate_coverletter_and_reasoning(job_title, job_desc, provider=provider, model_name=model)
                st.session_state.generated_coverletter = cover_result
                
                st.success("Documents generated!")
        else:
            st.error("Please provide Job Title and Description.")

with col2:
    st.header("2. Review & Export")
    
    tab1, tab2 = st.tabs(["📄 Resume", "✉️ Cover Letter"])
    
    with tab1:
        if st.session_state.generated_resume:
            resume_data = st.session_state.generated_resume.get('document', {})
            reasoning = st.session_state.generated_resume.get('reasoning', '')
            
            with st.expander("💡 AI Reasoning"):
                st.info(reasoning)
            
            # Editable Fields
            if isinstance(resume_data, dict):
                summary = st.text_area("Professional Summary", value=resume_data.get('summary', ''), height=100)
                skills = st.text_area("Key Skills", value=resume_data.get('skills', ''), height=100)
                experience = st.text_area("Relevant Experience", value=resume_data.get('experience', ''), height=200)
                projects = st.text_area("Key Projects", value=resume_data.get('projects', ''), height=150)
                
                if st.button("📥 Download Resume (DOCX)", use_container_width=True):
                    try:
                        template_path = os.path.join(DATA_DIR, "Manu_Martin_Resume_Template1.docx")
                        
                        # Create output directory
                        safe_company = re.sub(r'[\\/*?:"<>|]', "", company_name)
                        safe_job = re.sub(r'[\\/*?:"<>|]', "", job_title)
                        final_dir = os.path.join(OUTPUTS_DIR, safe_company, safe_job)
                        os.makedirs(final_dir, exist_ok=True)
                        
                        final_docx = os.path.join(final_dir, "Resume.docx")
                        
                        doc = docx.Document(template_path)
                        
                        # Map fields to template placeholders
                        replacements = {
                            '{{SUMMARY}}': summary,
                            '{{SKILLS}}': skills,
                            '{{EXPERIENCE}}': experience,
                            '{{PROJECTS}}': projects
                        }
                        
                        for placeholder, content in replacements.items():
                            replace_text_in_document(doc, placeholder, content)
                            
                        doc.save(final_docx)
                        
                        st.success("Resume saved!")
                        with open(final_docx, "rb") as f:
                            st.download_button("Download Resume DOCX", f.read(), file_name="Resume.docx")
                            
                    except Exception as e:
                        st.error(f"Error creating DOCX: {e}")
            else:
                st.error("Invalid Resume Data Format")

    with tab2:
        if st.session_state.generated_coverletter:
            cl_content = st.session_state.generated_coverletter.get('document', '')
            reasoning = st.session_state.generated_coverletter.get('reasoning', '')
            
            with st.expander("💡 AI Reasoning"):
                st.info(reasoning)
                
            cl_edited = st.text_area("Cover Letter Content", value=cl_content, height=400)
            
            if st.button("📥 Download Cover Letter (DOCX)", use_container_width=True):
                try:
                    template_path = os.path.join(DATA_DIR, "Manu_Martin_Coverletter_Template.docx")
                    
                    safe_company = re.sub(r'[\\/*?:"<>|]', "", company_name)
                    safe_job = re.sub(r'[\\/*?:"<>|]', "", job_title)
                    final_dir = os.path.join(OUTPUTS_DIR, safe_company, safe_job)
                    os.makedirs(final_dir, exist_ok=True)
                    
                    final_docx = os.path.join(final_dir, "CoverLetter.docx")
                    
                    doc = docx.Document(template_path)
                    
                    replace_text_in_document(doc, '{{COVER_LETTER_CONTENT}}', cl_edited)
                    
                    doc.save(final_docx)
                    
                    st.success("Cover Letter saved!")
                    with open(final_docx, "rb") as f:
                        st.download_button("Download Cover Letter DOCX", f.read(), file_name="CoverLetter.docx")
                        
                except Exception as e:
                    st.error(f"Error creating DOCX: {e}")
