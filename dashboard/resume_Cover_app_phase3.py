# dashboard/resume_Cover_app_phase3.py

import streamlit as st
import os
import re
import docx
from docx2pdf import convert
import pythoncom

# --- Setup Paths ---
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(_SCRIPT_DIR)
OUTPUTS_DIR = os.path.join(_PROJECT_ROOT, "outputs")
DATA_DIR = os.path.join(_PROJECT_ROOT, "data")

# --- Import your new AI function ---
from core.profile_memory_resume_phase2 import generate_document_and_reasoning

# --- Robust Template and Parsing Functions ---

# def replace_text_in_document(doc, placeholder, content):
#     """Finds and replaces a placeholder in all paragraphs and table cells."""
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for para in cell.paragraphs:
#                     if placeholder in para.text:
#                         para.text = para.text.replace(placeholder, content)
#     for para in doc.paragraphs:
#         if placeholder in para.text:
#             para.text = para.text.replace(placeholder, content)

def replace_text_in_document(doc, placeholder, content):
    """
    Finds and replaces a placeholder, correctly applying bold formatting
    for any text surrounded by double asterisks (e.g., **text**).
    """
    # This helper function does the complex work of adding formatted text
    def add_formatted_text(paragraph, text):
        # Split the text by the bold markers, keeping the markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # If the part is marked for bold, add it as a bold run
                # and remove the asterisks
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Otherwise, add it as a normal run
                run = paragraph.add_run(part)

    # --- Main logic to search the whole document ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        # Clear the placeholder text
                        para.text = ""
                        # Add the new, formatted content line by line
                        for line in content.split('\n'):
                            add_formatted_text(para, line)
                            para.add_run('\n') # Add line breaks back

    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = ""
            for line in content.split('\n'):
                add_formatted_text(para, line)
                para.add_run('\n')


def parse_ai_response_for_template_stars(text: str) -> dict:
    """Parses the AI's markdown-style output into a dictionary for the template."""
    sections = {}
    headers = ["Professional Summary", "Key Skills", "Relevant Projects", "Key Projects", "Relevant Experience", "Experience"]
    pattern = r"\*\*((" + "|".join(headers) + r"))\*\*"
    parts = re.split(pattern, text)
    for i in range(1, len(parts), 3):
        header = parts[i].strip()
        content = parts[i+2].strip()
        if "summary" in header.lower(): sections['{{SUMMARY}}'] = content
        elif "skills" in header.lower(): sections['{{SKILLS}}'] = content
        elif "projects" in header.lower(): sections['{{PROJECTS}}'] = content
        elif "experience" in header.lower(): sections['{{EXPERIENCE}}'] = content
    return sections

def parse_ai_response_for_template_hastag(text: str) -> dict:
    """Parses the AI's markdown-style output into a dictionary for the template."""
    sections = {}
    headers = ["Professional Summary", "Key Skills", "Relevant Projects", "Key Projects", "Relevant Experience", "Experience"]
    pattern = r"\#\#\# ((" + "|".join(headers) + r"))"
    # print(f"pattern: {pattern}")
    parts = re.split(pattern, text)
    # print(f"part: {parts}")
    # print(f"len(parts): {len(parts)}")
    for i in range(1, len(parts), 3):
        header = parts[i].strip()
        content = parts[i+2].strip()
        if "summary" in header.lower(): sections['{{SUMMARY}}'] = content
        elif "skills" in header.lower(): sections['{{SKILLS}}'] = content
        elif "projects" in header.lower(): sections['{{PROJECTS}}'] = content
        elif "experience" in header.lower(): sections['{{EXPERIENCE}}'] = content
    return sections



# --- Streamlit App ---
st.set_page_config(layout="wide")
st.title("🤖 Eva - AI Job Butler")

# --- UI for Job Input ---
st.header("1. Job Details")
job_title = st.text_input("Job Title", "AI Engineer")
company_name = st.text_input("Company Name", "Tech Innovations Inc.")
job_desc_input = st.text_area("Paste Full Job Description", height=200)

if st.button("✨ Generate Draft"):
    if job_desc_input and job_title:
        with st.spinner("Eva is reasoning and crafting your draft..."):
            response_dict = generate_document_and_reasoning(job_title, job_desc_input, "resume")
            st.session_state.document_draft = response_dict.get('document', '')
            st.session_state.reasoning = response_dict.get('reasoning', '')
    else:
        st.warning("Please provide a Job Title and Description.")

# --- UI for Editing and Approval ---
if 'document_draft' in st.session_state and st.session_state.document_draft:
    st.header("2. Review and Edit Draft")
    edited_text = st.text_area("Document Editor", st.session_state.document_draft, height=400)
    
    with st.expander("Show Eva's Reasoning 💡"):
        st.info(st.session_state.reasoning)

    st.header("3. Finalize and Export")
    if st.button("✅ Approve and Generate Final PDF"):
        try:
            # --- File Paths ---
            template_path = os.path.join(DATA_DIR, "Manu_Martin_Resume_Template1.docx")
            company_folder = re.sub(r'[\\/*?:"<>|]', "", company_name)
            job_folder = re.sub(r'[\\/*?:"<>|]', "", job_title)
            final_dir = os.path.join(OUTPUTS_DIR, company_folder, job_folder)
            os.makedirs(final_dir, exist_ok=True)
            temp_docx_path = os.path.join(final_dir, "temp.docx")
            final_pdf_path = os.path.join(final_dir, f"final_resume.pdf")

            # --- Fill Template and Convert ---
            doc = docx.Document(template_path)
            sections_to_fill = parse_ai_response_for_template_hastag(edited_text)
            # print(f"edited_text: {edited_text}")                  # For debugging
            # print(f"Sections to fill: {sections_to_fill}")        # For debugging
            if not sections_to_fill:
                sections_to_fill = parse_ai_response_for_template_stars(edited_text)
                print(f"Sections to fill after stars parsing: {sections_to_fill}")

            for placeholder, content in sections_to_fill.items():
                replace_text_in_document(doc, placeholder, content)

            doc.save(temp_docx_path)
            
            pythoncom.CoInitialize()
            convert(temp_docx_path, final_pdf_path)
            os.remove(temp_docx_path)
            
            st.success(f"Final PDF saved to: {final_pdf_path}")
            with open(final_pdf_path, "rb") as pdf_file:
                st.download_button("Download Final PDF", pdf_file.read(), file_name=f"final_resume.pdf")
                
        except Exception as e:
            st.error(f"Failed to generate document. Error: {e}")
        finally:
            pythoncom.CoUninitialize()