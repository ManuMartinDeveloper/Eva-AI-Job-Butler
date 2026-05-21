# backend/routers/generate.py
import os
import sys
import re
import docx
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, Dict

# --- Setup Paths ---
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.dirname(os.path.dirname(_SCRIPT_DIR))
sys.path.append(_PROJECT_ROOT)

from core.profile_memory_resume_phase2 import (
    generate_resume_and_reasoning,
    generate_coverletter_and_reasoning
)

router = APIRouter()

DATA_DIR = os.path.join(_PROJECT_ROOT, "data")
OUTPUTS_DIR = os.path.join(_PROJECT_ROOT, "outputs")

# Pydantic schemas
class GenerateRequest(BaseModel):
    job_title: str
    job_desc: str
    provider: Optional[str] = "gemini"
    model_name: Optional[str] = None

class ExportRequest(BaseModel):
    doc_type: str  # "resume" or "coverletter"
    company_name: str
    job_title: str
    content: Dict[str, str]  # Contains summary, skills, experience, projects, or text content

def replace_text_in_document(doc, placeholder, content):
    """Finds and replaces a placeholder, correctly applying bold formatting."""
    def add_formatted_text(paragraph, text):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                run = paragraph.add_run(part)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        para.text = ""
                        for line in content.split('\n'):
                            add_formatted_text(para, line)
                            para.add_run('\n')

    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = ""
            for line in content.split('\n'):
                add_formatted_text(para, line)
                para.add_run('\n')

@router.post("/resume")
def generate_resume(req: GenerateRequest):
    try:
        res = generate_resume_and_reasoning(
            job_title=req.job_title,
            job_desc=req.job_desc,
            provider=req.provider,
            model_name=req.model_name
        )
        return res
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failed: {str(e)}")

@router.post("/coverletter")
def generate_coverletter(req: GenerateRequest):
    try:
        res = generate_coverletter_and_reasoning(
            job_title=req.job_title,
            job_desc=req.job_desc,
            provider=req.provider,
            model_name=req.model_name
        )
        return res
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failed: {str(e)}")

@router.post("/export")
def export_docx(req: ExportRequest):
    safe_company = re.sub(r'[\\/*?:"<>|]', "", req.company_name)
    safe_job = re.sub(r'[\\/*?:"<>|]', "", req.job_title)
    
    final_dir = os.path.join(OUTPUTS_DIR, safe_company, safe_job)
    os.makedirs(final_dir, exist_ok=True)
    
    if req.doc_type == "resume":
        template_path = os.path.join(DATA_DIR, "Manu_Martin_Resume_Template1.docx")
        final_docx = os.path.join(final_dir, "Resume.docx")
        
        if not os.path.exists(template_path):
            raise HTTPException(status_code=404, detail="Resume template file not found.")
            
        try:
            doc = docx.Document(template_path)
            replacements = {
                '{{SUMMARY}}': req.content.get("summary", ""),
                '{{SKILLS}}': req.content.get("skills", ""),
                '{{EXPERIENCE}}': req.content.get("experience", ""),
                '{{PROJECTS}}': req.content.get("projects", "")
            }
            for placeholder, content in replacements.items():
                replace_text_in_document(doc, placeholder, content)
                
            doc.save(final_docx)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to create resume DOCX: {e}")
            
    elif req.doc_type == "coverletter":
        template_path = os.path.join(DATA_DIR, "Manu_Martin_coverletter_Template.docx")
        final_docx = os.path.join(final_dir, "CoverLetter.docx")
        
        if not os.path.exists(template_path):
            raise HTTPException(status_code=404, detail="Cover letter template file not found.")
            
        try:
            doc = docx.Document(template_path)
            replace_text_in_document(doc, '{{COVER_LETTER_CONTENT}}', req.content.get("text", ""))
            doc.save(final_docx)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to create cover letter DOCX: {e}")
    else:
        raise HTTPException(status_code=400, detail="Invalid doc_type. Choose 'resume' or 'coverletter'")

    # Return the relative path for download
    relative_path = os.path.relpath(final_docx, OUTPUTS_DIR)
    return {"file_path": relative_path, "filename": os.path.basename(final_docx)}

@router.get("/download")
def download_file(file_path: str):
    # Sanitize path to prevent directory traversal
    clean_path = os.path.abspath(os.path.join(OUTPUTS_DIR, file_path))
    if not clean_path.startswith(os.path.abspath(OUTPUTS_DIR)):
        raise HTTPException(status_code=403, detail="Access denied")
        
    if not os.path.exists(clean_path):
        raise HTTPException(status_code=404, detail="File not found")
        
    return FileResponse(
        clean_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(clean_path)
    )
